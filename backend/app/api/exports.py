from fastapi import APIRouter, HTTPException
from fastapi.responses import Response
from pydantic import BaseModel
from app.core.template_engine import template_engine
from playwright.async_api import async_playwright
import docx
from io import BytesIO

router = APIRouter()

class ExportRequest(BaseModel):
    cv_data: dict
    template_name: str = "professional-modern-v1"

@router.post("/preview")
async def preview_html(request: ExportRequest):
    try:
        html_content = template_engine.render_cv_html(request.cv_data, request.template_name)
        return Response(content=html_content, media_type="text/html")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/pdf")
async def export_pdf(request: ExportRequest):
    try:
        html_content = template_engine.render_cv_html(request.cv_data, request.template_name)
        
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            page = await browser.new_page()
            await page.set_content(html_content)
            # Remove margins and embed exact styling
            pdf_bytes = await page.pdf(print_background=True, format="A4", margin={"top": "0", "right": "0", "bottom": "0", "left": "0"})
            await browser.close()
            
        return Response(content=pdf_bytes, media_type="application/pdf", headers={"Content-Disposition": "attachment; filename=cv.pdf"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/docx")
async def export_docx(request: ExportRequest):
    try:
        # A simple DOCX generation based on the data
        doc = docx.Document()
        cv = request.cv_data
        
        # Name
        doc.add_heading(cv.get('personal_info', {}).get('name', 'Candidato'), 0)
        
        # Contact
        contact = f"{cv.get('personal_info', {}).get('email', '')} | {cv.get('personal_info', {}).get('phone', '')}"
        if cv.get('personal_info', {}).get('linkedin'):
             contact += f" | {cv.get('personal_info', {}).get('linkedin', '')}"
        doc.add_paragraph(contact)
        
        if cv.get('personal_info', {}).get('summary'):
            doc.add_paragraph(cv['personal_info']['summary'])
            
        doc.add_heading('Experience', level=1)
        for exp in cv.get('experience', []):
            end_date = "Present" if exp.get('present') else exp.get('end_date', '')
            p = doc.add_paragraph()
            p.add_run(f"{exp.get('role', '')} at {exp.get('company', '')}").bold = True
            p.add_run(f" ({exp.get('start_date', '')} - {end_date})").italic = True
            for point in exp.get('achievements', []):
                doc.add_paragraph(point, style='List Bullet')
                
        doc.add_heading('Education', level=1)
        for ed in cv.get('education', []):
            end_date = "Present" if ed.get('present') else ed.get('end_date', '')
            p = doc.add_paragraph()
            p.add_run(f"{ed.get('institution', '')}").bold = True
            p.add_run(f" - {ed.get('degree', '')} ({ed.get('start_date', '')} - {end_date})")
            
        doc.add_heading('Skills', level=1)
        skills_text = ", ".join(cv.get('skills', {}).get('technical', []))
        doc.add_paragraph(skills_text)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=cv.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DOCX Generation failed: {str(e)}")
